from datetime import datetime
from utils import validar_data
from storage import carregar_relatorio
from docx import Document
import os


def filtrar_por_intervalo(dados, data_inicio, data_fim):
    data_inicio_dt = datetime.strptime(data_inicio, "%d/%m/%Y")
    data_fim_dt = datetime.strptime(data_fim, "%d/%m/%Y")

    return [
        item for item in dados
        if data_inicio_dt <= datetime.strptime(item["data"], "%d/%m/%Y") <= data_fim_dt
    ]

from tkinter import messagebox

def aplicar_filtro(data_inicio, data_fim):
    if not validar_data(data_inicio) or not validar_data(data_fim):
        from tkinter import messagebox
        messagebox.showerror("Erro", "Datas inválidas!")
        return None

    dados = carregar_relatorio()
    filtrados = filtrar_por_intervalo(dados, data_inicio, data_fim)

    return filtrados


from docx import Document

def exportar_para_word(registros):
    if not registros:
        return

    # 📁 cria pasta
    pasta = "relatorios"
    os.makedirs(pasta, exist_ok=True)

    # 🕒 nome com timestamp
    nome_arquivo = f"relatorio_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"

    # 📄 caminho completo
    caminho = os.path.join(pasta, nome_arquivo)

    # 🧾 cria documento
    doc = Document()
    doc.add_heading("Relatório Filtrado", 0)

    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = 'Nome'
    table.rows[0].cells[1].text = 'Data'

    for item in registros:
        row = table.add_row().cells
        row[0].text = item["nome_arquivo"].replace(".docx", "")
        row[1].text = item["data"]

    # 🔥 salva no caminho correto
    doc.save(caminho)

    print(f"Arquivo salvo em: {os.path.abspath(caminho)}")